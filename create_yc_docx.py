import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Document Title
title_p = doc.add_paragraph()
title_run = title_p.add_run("Y Combinator Startup Internship Notes")
title_run.font.name = "Arial"
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(15, 23, 42)
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle_p = doc.add_paragraph()
sub_run = subtitle_p.add_run("Three Personalized Application Notes to YC Startups | Candidate: Saurabh Kumar")
sub_run.font.name = "Arial"
sub_run.font.size = Pt(12)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(71, 85, 105)
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# --- NOTE 1: LINEAR ---
doc.add_heading("1. Personalized Note to Linear (YC W19)", level=1)

p = doc.add_paragraph()
r = p.add_run("Target Startup: ")
r.bold = True
p.add_run("Linear (YC W19) | Focus: Issue Tracking & Product Operations")

p = doc.add_paragraph()
r = p.add_run("Skillset & Culture Fit: ")
r.bold = True
p.add_run("Linear sets the industry standard for fast, keyboard-driven, fluid user interfaces. My experience building real-time Kanban boards with Dragula JS, Tailwind CSS, and state-reconciled drag-and-drop aligns directly with Linear's frontend design principles.")

doc.add_paragraph().paragraph_format.space_after = Pt(4)

p_note = doc.add_paragraph()
p_note.paragraph_format.left_indent = Inches(0.3)
r_box = p_note.add_run("""Personalized Note:
"Hi Linear Engineering Team,

I'm a huge fan of Linear’s obsession with fluid user interaction and lightning-fast developer tools. Recently, I engineered and deployed Shiptivitas—a full-stack logistics productivity board featuring real-time Dragula drag-and-drop task reordering, status color transitions, and an Express/SQLite backend that persists task priority ranks dynamically.

I’ve deployed the full-stack app live on Render (https://shiptivitas-fullstack.onrender.com) and published the codebase to GitHub (https://github.com/SaurabhForge/shiptivitas-fullstack). I love building high-density, responsive web UIs with zero friction, and I’d be thrilled to contribute to Linear’s frontend and core application team as a Software Engineering Intern."
""")
r_box.font.name = "Calibri"
r_box.font.size = Pt(11)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# --- NOTE 2: VERCEL ---
doc.add_heading("2. Personalized Note to Vercel (YC S15)", level=1)

p = doc.add_paragraph()
r = p.add_run("Target Startup: ")
r.bold = True
p.add_run("Vercel (YC S15) | Focus: Frontend Cloud & Developer Infrastructure")

p = doc.add_paragraph()
r = p.add_run("Skillset & Culture Fit: ")
r.bold = True
p.add_run("Vercel empowers developers to ship web applications instantly with zero-config deployment. My experience building React applications, designing Stitch AI design tokens, and setting up automated Render blueprint deployments mirrors Vercel's mission.")

doc.add_paragraph().paragraph_format.space_after = Pt(4)

p_note2 = doc.add_paragraph()
p_note2.paragraph_format.left_indent = Inches(0.3)
r_box2 = p_note2.add_run("""Personalized Note:
"Hi Vercel Team,

As someone who cares deeply about web design, UI performance, and developer experience, I’ve followed Vercel’s impact on modern frontend engineering closely. Building full-stack web applications with React, Express, and Tailwind CSS has shown me how critical fast feedback loops are for product teams.

For my recent Shiptivitas project, I designed a minimal SaaS dashboard using Google Stitch design tokens, connected React components to a Node.js REST API with SQLite persistence, and configured automated Render deployment blueprints. You can view the live project at https://shiptivitas-fullstack.onrender.com. I would love to bring my passion for frontend craftsmanship and rapid prototyping to Vercel as a Software Engineering Intern."
""")
r_box2.font.name = "Calibri"
r_box2.font.size = Pt(11)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# --- NOTE 3: SUPABASE ---
doc.add_heading("3. Personalized Note to Supabase (YC S20)", level=1)

p = doc.add_paragraph()
r = p.add_run("Target Startup: ")
r.bold = True
p.add_run("Supabase (YC S20) | Focus: Open Source Backend Infrastructure & Relational Databases")

p = doc.add_paragraph()
r = p.add_run("Skillset & Culture Fit: ")
r.bold = True
p.add_run("Supabase makes relational databases accessible, reactive, and developer-friendly. In Task 2, I implemented relational SQLite database query handlers in Node.js to manage real-time card status moves and priority rank reordering across swimlanes.")

doc.add_paragraph().paragraph_format.space_after = Pt(4)

p_note3 = doc.add_paragraph()
p_note3.paragraph_format.left_indent = Inches(0.3)
r_box3 = p_note3.add_run("""Personalized Note:
"Hi Supabase Team,

I love Supabase’s mission to give developers real-time database superpowers without complex backend boilerplate. In my recent project, Shiptivitas, I implemented relational SQLite database query handlers in Node.js to manage real-time card status moves and priority rank reordering across swimlanes.

I enjoy working across the entire stack—from writing clean SQL queries to building interactive React frontends and deploying web services. Check out my backend repo at https://github.com/SaurabhForge/shiptivitas-backend-task2 and live demo at https://shiptivitas-fullstack.onrender.com. I’d love the opportunity to contribute to Supabase’s developer tooling as an Engineering Intern."
""")
r_box3.font.name = "Calibri"
r_box3.font.size = Pt(11)

docx_path = r'C:\Users\Saurabh Kumar\Desktop\shiptivitas-1\YC_Personalized_Notes.docx'
doc.save(docx_path)
print(f"Successfully generated {docx_path}")
